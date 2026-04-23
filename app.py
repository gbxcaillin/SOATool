"""
SOA Completion Agent — Python Backend
GBX Professional Services / Brightday Australia

Reads:  Fact Finder (.xlsx)  →  Fact Finder tab
        SOA Template (.docx) →  find & replace {{codes}}

Outputs: Completed SOA draft (.docx) with all insertions in red font.
         Unmapped codes are left as raw {{code}} text.

Run:
    pip install flask python-docx openpyxl
    python app.py
    Open http://localhost:5000
"""

from flask import Flask, request, send_file, jsonify, session, redirect, url_for, render_template_string
import openpyxl
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date
import io
import traceback
import re
import copy
import os
import hashlib

app = Flask(__name__)

# ─────────────────────────────────────────────
# AUTH CONFIG
# Read credentials from environment variables.
# Set these in Render dashboard — never hardcode.
# ─────────────────────────────────────────────
app.secret_key = os.environ.get("SECRET_KEY", "change-this-in-production")

# USERS dict — username: hashed password
# To generate a hash: python3 -c "import hashlib; print(hashlib.sha256('yourpassword'.encode()).hexdigest())"
# Add as many users as needed in the USERS env var format:
#   USERS=username1:hash1,username2:hash2
def load_users():
    users_env = os.environ.get("USERS", "")
    users = {}
    for entry in users_env.split(","):
        entry = entry.strip()
        if ":" in entry:
            username, pw_hash = entry.split(":", 1)
            users[username.strip().lower()] = pw_hash.strip()
    return users

def check_password(username, password):
    users = load_users()
    pw_hash = hashlib.sha256(password.encode()).hexdigest()
    return users.get(username.lower()) == pw_hash

def logged_in():
    return session.get("authenticated") is True

# ─────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────
RED   = RGBColor(0xFF, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)

# Fact Finder tab — column letters for multi-fund fields (up to 5 funds)
# Columns: B=2, D=4, F=6, H=8, J=10
FUND_COLS = [2, 4, 6, 8, 10]


# ─────────────────────────────────────────────
# FACT FINDER READER
# ─────────────────────────────────────────────

def read_fact_finder(xlsx_bytes, risk_profile, no_insurance_flag):
    """
    Read the Fact Finder xlsx and return:
        - data dict  { "{{CODE}}": "value" }
        - conditionals dict  { "DELETE_KEY": True/False }
    """
    wb = openpyxl.load_workbook(io.BytesIO(xlsx_bytes), data_only=True)
    ws = wb["Fact Finder"]

    def cell(row, col):
        """Return cleaned string value from a cell, or '' if empty/zero placeholder."""
        v = ws.cell(row=row, column=col).value
        if v is None:
            return ""
        s = str(v).strip()
        # The FF uses 0 as a placeholder for unfilled cells
        if s in ("0", "00:00:00", "#REF!"):
            return ""
        return s

    def cells_across(row, cols=FUND_COLS):
        """Return list of non-empty values across multiple fund columns."""
        return [cell(row, c) for c in cols if cell(row, c)]

    def join_funds(row, sep=", "):
        vals = cells_across(row)
        return sep.join(vals) if vals else ""

    def sum_funds(row):
        total = 0
        for c in FUND_COLS:
            v = ws.cell(row=row, column=c).value
            try:
                total += float(str(v).replace(",","").replace("$",""))
            except Exception:
                pass
        return total

    def currency(row, col):
        v = ws.cell(row=row, column=col).value
        try:
            return f"${float(str(v).replace(',','').replace('$','')):,.0f}"
        except Exception:
            return ""

    def currency_sum(row):
        s = sum_funds(row)
        return f"${s:,.0f}" if s else ""

    def age_from_dob(row, col=2):
        """Calculate age from a date cell."""
        v = ws.cell(row=row, column=col).value
        if not v:
            return None
        try:
            from datetime import datetime
            if hasattr(v, 'year'):
                dob = v
            else:
                for fmt in ("%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y"):
                    try:
                        dob = datetime.strptime(str(v), fmt)
                        break
                    except ValueError:
                        continue
                else:
                    return None
            today = date.today()
            return today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))
        except Exception:
            return None

    def format_date(row, col=2):
        v = ws.cell(row=row, column=col).value
        if not v:
            return ""
        try:
            if hasattr(v, 'strftime'):
                return v.strftime("%d/%m/%Y")
            return str(v)
        except Exception:
            return str(v)

    # ── Personal Details ──
    title       = cell(10, 2)
    first_name  = cell(11, 2)
    middle_name = cell(12, 2)
    last_name   = cell(13, 2)
    full_name_parts = [p for p in [first_name, middle_name, last_name] if p]
    full_name   = " ".join(full_name_parts)
    dob_str     = format_date(15, 2)
    age         = age_from_dob(15, 2)
    phone       = cell(16, 2)
    email       = cell(17, 2)
    address_parts = [p for p in [cell(18,2), cell(19,2), cell(20,2), cell(21,2)] if p]
    address     = ", ".join(address_parts)

    # ── Employment ──
    occupation  = cell(28, 2)
    emp_status  = cell(23, 2)

    # ── Income ──
    gross_income_raw = ws.cell(row=32, column=2).value
    try:
        gross_income_num = float(str(gross_income_raw).replace(",","").replace("$",""))
        gross_income = f"${gross_income_num:,.0f}"
    except Exception:
        gross_income_num = 0
        gross_income = ""

    sgc_pct_raw = ws.cell(row=34, column=2).value
    try:
        sgc_pct = float(str(sgc_pct_raw).replace("%","")) / 100
    except Exception:
        sgc_pct = 0.12
    super_contribution = f"${gross_income_num * sgc_pct:,.0f}" if gross_income_num else ""

    salary_sacrifice_raw = ws.cell(row=35, column=2).value
    try:
        salary_sacrifice = f"${float(str(salary_sacrifice_raw).replace(',','').replace('$','')):,.0f}"
        annualised_salary_sacrifice = salary_sacrifice
    except Exception:
        salary_sacrifice = ""
        annualised_salary_sacrifice = ""

    # ── Retirement Age ──
    ret_age_1 = cell(8, 2)
    ret_age_2 = cell(9, 2)
    retirement_age = ret_age_2 if ret_age_2 else ret_age_1

    # ── Spouse ──
    spouse_dob  = format_date(47, 2)
    spouse_income_raw = ws.cell(row=49, column=2).value
    try:
        spouse_income = f"${float(str(spouse_income_raw).replace(',','').replace('$','')):,.0f}"
    except Exception:
        spouse_income = ""
    spouse_balance_raw = ws.cell(row=50, column=2).value
    try:
        spouse_balance = f"${float(str(spouse_balance_raw).replace(',','').replace('$','')):,.0f}"
    except Exception:
        spouse_balance = ""

    # ── Dependants ──
    has_spouse = bool(cell(46, 2))  # Spouse Name row
    dep_ages = cells_across(56)
    no_dependants = (1 if has_spouse else 0) + len(dep_ages)

    # ── Assets & Liabilities ──
    primary_residence_val = currency(73, 2)
    primary_residence_debt = currency(74, 2)
    investment_prop_val = currency(76, 2)
    investment_prop_debt = currency(77, 2)
    other_asset1_val = currency(79, 2)
    personal_loan1_val = currency(81, 2)

    # Total assets
    total_assets = 0
    for r, c in [(73,2),(76,2),(79,2)]:
        v = ws.cell(row=r, column=c).value
        try:
            total_assets += float(str(v).replace(",","").replace("$",""))
        except Exception:
            pass
    # Add super balances
    total_assets += sum_funds(94)
    total_assets_str = f"${total_assets:,.0f}" if total_assets else ""

    total_liabilities = 0
    for r, c in [(74,2),(77,2),(81,2)]:
        v = ws.cell(row=r, column=c).value
        try:
            total_liabilities += float(str(v).replace(",","").replace("$",""))
        except Exception:
            pass
    total_liabilities_str = f"${total_liabilities:,.0f}" if total_liabilities else ""

    # ── Super Funds ──
    current_super_funds = join_funds(92)
    current_super_balance = currency_sum(94)
    current_balance = current_super_balance

    # ── Insurance across funds ──
    def insurance_across(row):
        vals = []
        for c in FUND_COLS:
            v = ws.cell(row=row, column=c).value
            if v:
                s = str(v).strip()
                if s not in ("0","","None"):
                    vals.append(s)
        if not vals:
            return ""
        # Format: if multiple funds, join; if same value, show once
        unique = list(dict.fromkeys(vals))
        return " / ".join(unique)

    life_ins  = insurance_across(102)
    tpd_ins   = insurance_across(103)
    ip_month  = insurance_across(104)
    ip_wait   = insurance_across(105)
    ip_benefit = insurance_across(106)
    premiums  = insurance_across(107)

    # ── Binding Death Nominee ──
    nominee_names = []
    for c in [2, 4, 6, 8, 10]:
        v = ws.cell(row=62, column=c).value
        if v and str(v).strip() not in ("0","","None"):
            nominee_names.append(str(v).strip())
    binding_death_nominee = ", ".join(nominee_names) if nominee_names else "N/A"

    # ── Current Date ──
    current_date = date.today().strftime("%d %B %Y")

    # ── Risk Profile (from UI selection) ──
    current_risk_profile = risk_profile  # passed in from form

    # ─────────────────────────────────
    # BUILD DATA DICT
    # ─────────────────────────────────
    data = {
        "{{Title}}":                             title,
        "{{ClientFullName}}":                    full_name,
        "{{ClientFirstName}}":                   first_name,
        "{{ClientDOB}}":                         dob_str,
        "{{ClientAddress}}":                     address,
        "{{ClientPhone}}":                       phone,
        "{{ClientEmail}}":                       email,
        "{{ClientOccupation}}":                  occupation,
        "{{ClientSalary}}":                      gross_income,
        "{{fld_SuperContribution}}":             super_contribution,
        "{{fld_SalarySacrifice}}":               salary_sacrifice,
        "{{CurrentSuperFunds}}":                 current_super_funds,
        "{{SpouseDOB}}":                         spouse_dob,
        "{{SpouseIncome}}":                      spouse_income,
        "{{SpouseBalance}}":                     spouse_balance,
        "{{NoDependants}}":                      str(no_dependants),
        "{{fld_CurrentSuperannuationBalance}}":  current_super_balance,
        "{{CurrentLifeInsurance}}":              life_ins,
        "{{CurrentTPDInsurance}}":               tpd_ins,
        "{{CurrentIncomeProtectionPerMonth}}":   ip_month,
        "{{CurrentIncomeProtectionWaitingPeriod}}": ip_wait,
        "{{CurrentIncomeProtectionBenefitPeriod}}": ip_benefit,
        "{{CurrentSuperPremiums}}":              premiums,
        "{{ValueOfPrimaryResidence}}":           primary_residence_val,
        "{{DebtOnPrimaryResidence}}":            primary_residence_debt,
        "{{ValueOfInvestmentProperty}}":         investment_prop_val,
        "{{DebtOnInvestmentProperty}}":          investment_prop_debt,
        "{{OtherAsset1Value}}":                  other_asset1_val,
        "{{PersonalLoan1Value}}":                personal_loan1_val,
        "{{TotalAssetValue}}":                   total_assets_str,
        "{{TotalLiabilityValue}}":               total_liabilities_str,
        "{{RetirementAge}}":                     retirement_age,
        "{{CurrentBalance}}":                    current_balance,
        "{{CurrentAge}}":                        str(age) if age else "",
        "{{CurrentDate}}":                       current_date,
        "{{AnnualisedSalarySacrificeAmount}}":   annualised_salary_sacrifice,
        "{{BindingDeathNominee}}":               binding_death_nominee,
        "{{CurrentRiskProfile}}":                current_risk_profile,
        # Goals — left as raw codes (unmapped)
        # Table codes — left as raw codes (unmapped)
    }

    # ─────────────────────────────────
    # CONDITIONALS
    # ─────────────────────────────────
    total_balance = sum_funds(94)

    # Row 100: insurance in fund — check all fund columns
    has_any_insurance = any(
        str(ws.cell(row=100, column=c).value or "").strip().lower() == "yes"
        for c in FUND_COLS
    )

    # Row 108: medically underwritten
    has_underwritten = any(
        str(ws.cell(row=108, column=c).value or "").strip().lower() == "medically underwritten"
        for c in FUND_COLS
    )

    conditionals = {
        # True = DELETE this block
        "DeleteIfAgeGreaterThan55":              (age is not None and age >= 55),
        "DeleteIfAgeLessThan55":                 (age is not None and age < 55),
        "DeleteIfBalanceBelow500k":              (total_balance < 500_000),
        "DeleteIfNoCurrentInsurance":            (not has_any_insurance),
        "DeleteIfNoInsuranceAtAll":              no_insurance_flag,   # UI checkbox
        "DeleteIfNoScopedInsurance":             False,  # unmapped — never delete
        "DeleteIfNoScopedTrauma":                False,  # unmapped — never delete
        "DeleteIfNoTrauma":                      False,  # unmapped — never delete
        "DeleteIfPersonalDeductibleContributions": False,  # unmapped — never delete
        "DeleteifNoCurrentUnderwrittenInsurance": (not has_underwritten),
    }

    return data, conditionals


# ─────────────────────────────────────────────
# SOA DOCUMENT PROCESSOR
# ─────────────────────────────────────────────

# Codes that are intentionally left as raw {{code}} — never replaced
UNMAPPED_CODES = {
    "{{Date}}",
    "{{OtherAsset1}}",
    "{{OtherAsset2}}",
    "{{OtherAsset2Value}}",
    "{{PersonalLoan2Value}}",
    "{{****PersonalLoan1}}",
    "{{****PersonalLoan2}}",
    "{{NeedsAnalysisLifeInsurance}}",
    "{{NeedsAnalysisTPD}}",
    "{{NeedsAnalysisIP}}",
    "{{NeedsAnalysisTrauma}}",
    "{{Tbl_SalarySacrifice}}",
    "{{tbl_CurrentSuperFundsRiskProfilePerformance}}",
    "{{Make personal deductible contributions/Salary sacrifice}}",
    "{{DeleteIfNoScopedInsurance}}",
    "{{EndDeleteIfNoScopedInsurance}}",
    "{{DeleteIfNoScopedTrauma}}",
    "{{DeleteIfNoTrauma}}",
    "{{EndDeleteIfNoTrauma}}",
    "{{DeleteIfPersonalDeductibleContributions}}",
    "{{EndDeleteIfPersonalDeductibleContributions}}",
    "{{CurrentInsuer}}",
    "{{SalarySacrificeAmount}}",
    "{{SalarySacrificeFrequency}}",
    "{{NetTaxSavings}}",
    "{{zzz}}",
    "{{SuperGoal}}",
    "{{InsuranceGoal}}",
    "{{SalarySacrificeGoal}}",
    "{{EstatePlanningGoal}}",
    "{{RetirementGoal}}",
    "{{DeleteIfNoInsuranceAtAll}}",
    "{{EndDeleteIfNoInsuranceAtAll}}",
}

# Pair up conditional block tags
CONDITIONAL_PAIRS = [
    ("{{DeleteIfAgeGreaterThan55}}",              "{{EndDeleteIfAgeGreaterThan55}}",              "DeleteIfAgeGreaterThan55"),
    ("{{DeleteIfAgeLessThan55}}",                 "{{EndDeleteIfAgeLessThan55}}",                 "DeleteIfAgeLessThan55"),
    ("{{DeleteIfBalanceBelow500k}}",              "{{EndDeleteIfBalanceBelow500k}}",              "DeleteIfBalanceBelow500k"),
    ("{{DeleteIfNoCurrentInsurance}}",            "{{EndDeleteIfNoCurrentInsurance}}",            "DeleteIfNoCurrentInsurance"),
    ("{{DeleteifNoCurrentUnderwrittenInsurance}}","{{EndDeleteifNoCurrentUnderwrittenInsurance}}","DeleteifNoCurrentUnderwrittenInsurance"),
]

# For DeleteIfNoInsuranceAtAll — single tag (no end tag), marks start of section to delete
# We treat the content following it until next section heading as the block
NO_INSURANCE_SINGLE_TAG = "{{DeleteIfNoInsuranceAtAll}}"


def get_full_text(paragraph):
    return "".join(run.text for run in paragraph.runs)


def para_contains(paragraph, code):
    return code in get_full_text(paragraph)


def replace_code_in_run(run, code, value, use_red):
    """Replace a code in a single run, applying red font to the replacement."""
    if code not in run.text:
        return
    parts = run.text.split(code)
    # If only one part before and after — simple case
    if len(parts) == 2:
        before, after = parts
        run.text = before
        # Insert red replacement run after this run
        p = run._r.getparent()
        idx = list(p).index(run._r)

        def make_run(text, red):
            from docx.oxml import OxmlElement
            r_el = OxmlElement('w:r')
            # Copy rPr from original run
            if run._r.find(qn('w:rPr')) is not None:
                rPr = copy.deepcopy(run._r.find(qn('w:rPr')))
                # Set or remove colour
                color_el = rPr.find(qn('w:color'))
                if color_el is None:
                    color_el = OxmlElement('w:color')
                    rPr.append(color_el)
                if red:
                    color_el.set(qn('w:val'), 'FF0000')
                else:
                    color_el.set(qn('w:val'), 'auto')
                r_el.append(rPr)
            else:
                if red:
                    rPr = OxmlElement('w:rPr')
                    color_el = OxmlElement('w:color')
                    color_el.set(qn('w:val'), 'FF0000')
                    rPr.append(color_el)
                    r_el.append(rPr)
            t_el = OxmlElement('w:t')
            t_el.text = text
            if text.startswith(' ') or text.endswith(' '):
                t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            r_el.append(t_el)
            return r_el

        # Insert replacement
        if value:
            p.insert(idx + 1, make_run(value, use_red))
        # Insert after-text
        if after:
            p.insert(idx + 2, make_run(after, False))
    else:
        # Multiple occurrences in one run — replace all
        new_text = run.text.replace(code, value)
        run.text = new_text
        if use_red and value:
            run.font.color.rgb = RED


def process_paragraph_text(paragraph, data, unmapped):
    """Replace all known codes in a paragraph. Leave unmapped codes untouched."""
    full = get_full_text(paragraph)
    if "{{" not in full:
        return

    # Find all codes in this paragraph
    codes_present = re.findall(r'\{\{[^}]+\}\}', full)

    for code in codes_present:
        if code in unmapped:
            continue  # Leave raw
        if code in data:
            value = data[code]
            # Work run by run
            for run in paragraph.runs:
                if code in run.text:
                    replace_code_in_run(run, code, value, use_red=True)


def process_table(table, data, unmapped):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                process_paragraph_text(paragraph, data, unmapped)
            for nested_table in cell.tables:
                process_table(nested_table, data, unmapped)


def collect_all_paragraphs(doc):
    """Return flat list of (paragraph, parent_element, index) for body + tables."""
    items = []
    body = doc.element.body
    for i, child in enumerate(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            from docx.text.paragraph import Paragraph
            items.append(Paragraph(child, doc))
        elif tag == 'tbl':
            from docx.table import Table
            tbl = Table(child, doc)
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        items.append(p)
    return items


def apply_conditional_deletions(doc, conditionals):
    """
    Walk through document body elements.
    When a start-tag paragraph is found and its condition is True,
    collect and remove all elements up to and including the end-tag paragraph.
    """
    body = doc.element.body
    elements = list(body)

    def get_para_text(el):
        return "".join(t.text or "" for t in el.iter(qn('w:t')))

    for start_tag, end_tag, condition_key in CONDITIONAL_PAIRS:
        should_delete = conditionals.get(condition_key, False)
        if not should_delete:
            # Still remove the marker tags themselves (they're not content)
            to_remove = []
            for el in list(body):
                txt = get_para_text(el)
                if start_tag in txt or end_tag in txt:
                    to_remove.append(el)
            for el in to_remove:
                body.remove(el)
            continue

        # Delete everything between (and including) start and end tags
        in_block = False
        to_remove = []
        for el in list(body):
            txt = get_para_text(el)
            if start_tag in txt:
                in_block = True
            if in_block:
                to_remove.append(el)
            if end_tag in txt and in_block:
                in_block = False
        for el in to_remove:
            try:
                body.remove(el)
            except ValueError:
                pass

    # Handle DeleteIfNoInsuranceAtAll (no end tag)
    # Remove the single marker tag paragraph regardless
    should_delete_no_ins = conditionals.get("DeleteIfNoInsuranceAtAll", False)
    to_remove = []
    in_block = False
    for el in list(body):
        txt = get_para_text(el)
        if NO_INSURANCE_SINGLE_TAG in txt:
            to_remove.append(el)  # always remove the tag itself
            if should_delete_no_ins:
                in_block = True
            continue
        if in_block:
            # Delete until we hit the next heading-level paragraph or end of section
            # Heuristic: stop at next paragraph that has bold text > 12pt or is a heading style
            tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
            if tag == 'p':
                style = el.find('.//' + qn('w:pStyle'))
                style_val = style.get(qn('w:val'), '') if style is not None else ''
                if 'Heading' in style_val or style_val.startswith('h'):
                    in_block = False
                    continue
            to_remove.append(el)
    for el in to_remove:
        try:
            body.remove(el)
        except ValueError:
            pass


def process_soa(template_bytes, data, conditionals):
    """Main processor — returns completed docx as bytes."""
    doc = Document(io.BytesIO(template_bytes))

    # Step 1: Apply conditional block deletions
    apply_conditional_deletions(doc, conditionals)

    # Step 2: Replace codes in body paragraphs
    for paragraph in doc.paragraphs:
        process_paragraph_text(paragraph, data, UNMAPPED_CODES)

    # Step 3: Replace codes in tables
    for table in doc.tables:
        process_table(table, data, UNMAPPED_CODES)

    # Step 4: Replace codes in headers and footers
    for section in doc.sections:
        for hdr in [section.header, section.footer,
                    section.even_page_header, section.even_page_footer,
                    section.first_page_header, section.first_page_footer]:
            if hdr:
                for paragraph in hdr.paragraphs:
                    process_paragraph_text(paragraph, data, UNMAPPED_CODES)
                for table in hdr.tables:
                    process_table(table, data, UNMAPPED_CODES)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out


# ─────────────────────────────────────────────
# LOGIN PAGE HTML
# ─────────────────────────────────────────────

LOGIN_HTML = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>SOA Agent — Login</title>
<link href="https://fonts.googleapis.com/css2?family=Cormorant+Garamond:wght@300;400;600&family=Outfit:wght@300;400;500;600&display=swap" rel="stylesheet">
<style>
  :root {
    --navy: #0d1b2e; --navy-mid: #122440; --gold: #c9a84c;
    --gold-light: #e2c47a; --gold-dim: #8a6f2e;
    --white: #f5f3ee; --white-dim: #b8b4aa; --red: #e84040;
    --border: rgba(201,168,76,0.18);
  }
  *, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }
  body {
    background: var(--navy); color: var(--white);
    font-family: 'Outfit', sans-serif; font-weight: 300;
    min-height: 100vh; display: flex; align-items: center; justify-content: center;
  }
  body::before {
    content: ''; position: fixed; inset: 0;
    background-image: linear-gradient(rgba(201,168,76,0.03) 1px, transparent 1px),
      linear-gradient(90deg, rgba(201,168,76,0.03) 1px, transparent 1px);
    background-size: 48px 48px; pointer-events: none;
  }
  body::after {
    content: ''; position: fixed; top: -120px; right: -120px;
    width: 480px; height: 480px;
    background: radial-gradient(circle, rgba(201,168,76,0.08) 0%, transparent 70%);
    pointer-events: none;
  }
  .card {
    position: relative; z-index: 1;
    background: var(--navy-mid); border: 1px solid var(--border);
    padding: 48px 44px; width: 100%; max-width: 420px;
    animation: fadeUp 0.6s ease both;
  }
  .logo-area { display: flex; align-items: center; gap: 14px; margin-bottom: 40px; }
  .logo-mark {
    width: 42px; height: 42px; border: 1.5px solid var(--gold);
    display: flex; align-items: center; justify-content: center;
    font-family: 'Cormorant Garamond', serif; font-size: 18px;
    font-weight: 600; color: var(--gold); letter-spacing: 1px; flex-shrink: 0;
  }
  .logo-text .brand {
    font-family: 'Cormorant Garamond', serif; font-size: 15px; font-weight: 600;
    color: var(--white); letter-spacing: 1.5px; text-transform: uppercase; line-height: 1; display: block;
  }
  .logo-text .sub {
    font-size: 9px; color: var(--gold); letter-spacing: 2.5px;
    text-transform: uppercase; margin-top: 4px; display: block;
  }
  h1 {
    font-family: 'Cormorant Garamond', serif; font-size: 28px; font-weight: 300;
    line-height: 1.2; margin-bottom: 8px;
  }
  h1 em { font-style: italic; color: var(--gold-light); }
  .subtitle { font-size: 12px; color: var(--white-dim); margin-bottom: 36px; line-height: 1.6; }
  .field { margin-bottom: 20px; }
  .field label {
    display: block; font-size: 10px; font-weight: 500;
    letter-spacing: 2px; text-transform: uppercase; color: var(--white-dim); margin-bottom: 8px;
  }
  .field input {
    width: 100%; background: var(--navy); border: 1px solid var(--border);
    color: var(--white); font-family: 'Outfit', sans-serif; font-size: 13px;
    padding: 12px 14px; outline: none; transition: border-color 0.2s;
  }
  .field input:focus { border-color: var(--gold); }
  .btn-login {
    width: 100%; background: var(--gold); color: var(--navy); border: none;
    padding: 14px; font-family: 'Outfit', sans-serif; font-size: 11px;
    font-weight: 600; letter-spacing: 2.5px; text-transform: uppercase;
    cursor: pointer; margin-top: 8px; transition: opacity 0.2s;
  }
  .btn-login:hover { opacity: 0.88; }
  .error {
    background: rgba(232,64,64,0.08); border: 1px solid rgba(232,64,64,0.3);
    color: var(--red); font-size: 12px; padding: 10px 14px; margin-bottom: 20px;
  }
  .footer-note { font-size: 10px; color: var(--white-dim); opacity: 0.5; margin-top: 28px; text-align: center; letter-spacing: 0.5px; }
  @keyframes fadeUp { from { opacity: 0; transform: translateY(20px); } to { opacity: 1; transform: translateY(0); } }
</style>
</head>
<body>
<div class="card">
  <div class="logo-area">
    <div class="logo-mark">GBX</div>
    <div class="logo-text">
      <span class="brand">GBX Professional Services</span>
      <span class="sub">Financial Intelligence Suite</span>
    </div>
  </div>
  <h1>SOA <em>Completion</em><br>Agent</h1>
  <p class="subtitle">Internal access only. Enter your credentials to continue.</p>
  {% if error %}
  <div class="error">{{ error }}</div>
  {% endif %}
  <form method="POST" action="/login">
    <div class="field">
      <label>Username</label>
      <input type="text" name="username" autocomplete="username" autofocus required>
    </div>
    <div class="field">
      <label>Password</label>
      <input type="password" name="password" autocomplete="current-password" required>
    </div>
    <button class="btn-login" type="submit">Sign In →</button>
  </form>
  <p class="footer-note">GBX PS Pty Ltd · ABN 45 674 252 905 · Internal Use Only</p>
</div>
</body>
</html>"""


# ─────────────────────────────────────────────
# FLASK ROUTES
# ─────────────────────────────────────────────

@app.route("/login", methods=["GET", "POST"])
def login():
    error = None
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")
        if check_password(username, password):
            session["authenticated"] = True
            session["username"] = username.lower()
            return redirect(url_for("tool"))
        else:
            error = "Invalid username or password."
    return render_template_string(LOGIN_HTML, error=error)


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))


@app.route("/")
def tool():
    if not logged_in():
        return redirect(url_for("login"))
    with open("index.html", "r") as f:
        return f.read()


@app.route("/process", methods=["POST"])
def process():
    if not logged_in():
        return jsonify({"error": "Not authenticated"}), 401
    try:
        if "fact_finder" not in request.files:
            return jsonify({"error": "Missing Fact Finder file"}), 400
        if "soa_template" not in request.files:
            return jsonify({"error": "Missing SOA Template file"}), 400

        risk_profile = request.form.get("risk_profile", "").strip()
        no_insurance = request.form.get("no_insurance", "false").lower() == "true"

        if not risk_profile:
            return jsonify({"error": "Risk profile must be selected"}), 400

        ff_bytes       = request.files["fact_finder"].read()
        template_bytes = request.files["soa_template"].read()

        data, conditionals = read_fact_finder(ff_bytes, risk_profile, no_insurance)
        out = process_soa(template_bytes, data, conditionals)

        client_name = data.get("{{ClientFullName}}", "Client")
        today = date.today().strftime("%Y%m%d")
        filename = f"SOA_Draft_{client_name.replace(' ','_')}_{today}.docx"

        return send_file(
            out,
            as_attachment=True,
            download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except KeyError as e:
        return jsonify({"error": f"Fact Finder tab not found or unexpected structure: {e}"}), 400
    except Exception as e:
        return jsonify({"error": str(e), "trace": traceback.format_exc()}), 500


if __name__ == "__main__":
    app.run(debug=False, port=5000)
